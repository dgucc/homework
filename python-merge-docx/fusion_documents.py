from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import os
import copy

def merge_docx_files(file_paths, output_path):
    # Créer un nouveau document
    merged_document = Document()

    for file_path in file_paths:
        # Vérifier que le fichier existe
        if not os.path.exists(file_path):
            print(f"Fichier non trouvé : {file_path}")
            continue

        # Ajouter le chemin comme titre
        title_paragraph = merged_document.add_paragraph()
        run = title_paragraph.add_run(file_path)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)  # violet 
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p = title_paragraph._element
        
        # "Insert(-1, titleText) <=> append...
        merged_document._element.body.insert(-1, p)
        
        # Ouvrir le document source
        sub_doc = Document(file_path)

        # Copier chaque élément (paragraphes, tableaux, etc.) dans le document final
        for element in sub_doc.element.body:
            # Faire une copie profonde de l'élément pour éviter les problèmes de référence
            merged_document.element.body.append(copy.deepcopy(element))            


    # Sauvegarder le document fusionné
    merged_document.save(output_path)
    print(f"Document fusionné sauvegardé sous : {output_path}")

# === Utilisation ===
file_list = [
    "files/doc1.docx",
    "files/doc2.docx",
    "files/doc3.docx"
]

merge_docx_files(file_list, "result.docx")